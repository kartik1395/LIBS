from PIL import Image
import numpy as np
import cv2
import pandas as pd
import random
from extract import *
from openpyxl.drawing.image import Image as oim
from docx import Document
from docx.shared import Inches
from datetime import datetime
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH


incorrect = {}
correct = 0
incorrect_books = 0
df = pd.DataFrame(columns=['shelf', 'book'])
# global document 
# document= Document()
# document.add_heading('Misplaced Books Report')

def OCR(img):  # Where teerapong can add his text extraction code
    # assuming you return a single string per label ex - 'MAP-NAR'/'ABC'
    text = random.choice(["XCV", "ABC"])
    text_shelf = "ABC-ABD"
    return text, text_shelf


# img = Image.open('LIBRABYIMAGE')

def extractLabels(result, document):
    # labels_boxes, shelf_boxes
    # 15 and 2 are just arbitary numbers suggesting that 15 books and 2 shelf labels were detected
    # label_boxes = [[1]*4]*15  # A list of all bounding boxes of labels -> [[x1,y1,x2,y2],[x1,y1,x2,y2]..]
    # shelf_boxes = [[1]*4]*2  # A list of all bounding boxes of shelfs -> [[x1,y1,x2,y2],[x1,y1,x2,y2]..]

    text_label = []  # Creating a list of text labels to be appened

 #    document= Document()
	# document.add_heading('Misplaced Books Report')
    
    for item in result:
        label_boxes = item['labels']
        shelf_boxes = item['shelves']
        img_name = item['filename']
        img = Image.open(img_name)

        for i in range(0, len(shelf_boxes)):  # iterate over all shelf labels
            sx1, sy1, sx2, sy2 = shelf_boxes[i][0], shelf_boxes[i][1], shelf_boxes[i][2], shelf_boxes[i][3]

            # for each shelf label, check if the book label is to be matched with the shelf label below the book
            for j in range(0, len(label_boxes)):
                lx1, ly1, lx2, ly2 = label_boxes[j][0], label_boxes[j][1], label_boxes[j][2], label_boxes[j][3]
                roi_label = img.crop((lx1, ly1, lx2, ly2))
                # print("Shelf", sx1, sy1, sx2, sy2,
                #       "labels", lx1, ly1, lx2, ly2)
                # roi_label.save('test'+str(j)+'.jpg')
                # k = 1
                # compare the four coordinates of book label and shelf label
                if lx1 > sx1 and ly1 < sy1 and lx2 > sx2 and ly2 < sy2:

                    # crop the image for book label
                    roi_label = img.crop((lx1, ly1, lx2, ly2))
                    # crop the image for shelf label
                    roi_shelf = img.crop((sx1 -100 ,sy1, sx2, sy2))
                    # roi_label.save('test'+str(k)+'.jpg')
                    # k+=1
                    # text, _ = OCR(roi_label)
                    text = getText(roi_label, 'label')
                    text_shelf = getText(roi_shelf, 'shelf')

                    # _, text_shelf = OCR(roi_shelf)
                    text_label.append(text)
                    shelf_label = text_shelf

                    # text_label.append(OCR(roi_label)) #append the book label to the list after extracting label
                    # shelf_label = OCR(roi_shelf) # store the self label only if the shelf label is below book label

        
        createReport(text_label, shelf_label, df, img_name, document)
        

    # return text_label, shelf_label


def createReport(text_label, shelf_label, df, img_name, document):

    incorrect_books = 0
    total_inc = 0
    # writer = pd.ExcelWriter('test.xlsx', engine='openpyxl')
    # df.to_excel(writer, sheet_name = 'Report')

    misplaced = []

    for text in text_label:  # iterate over all book labels extracted

        shelf_class = shelf_label.split('-')
        # print("SC",shelf_class)
        if len(shelf_class) > 1:  # if the shelf label has a range

            if text >= shelf_class[0] and text <= shelf_class[1]:  # Correct book
                correct1 = correct + 1
                wrong = False
                # print("SC>1 correct",text)
            else:
                incorrect['shelf'] = shelf_label
                incorrect['book'] = text
                # print("SC>1 incorrect",text)
                incorrect_books += 1
                misplaced.append((shelf_label, text))
                

                wrong = True

                # push incorrect book label and the shelf it is kept on
                df = df.append(incorrect, ignore_index=True)

        else:  # if no range and needs exact match
            if text == shelf_class[0]:  # Correct
                correct1 = correct + 1
                wrong = False
                # print("SC=1 correct",text)
            else:
                incorrect['shelf'] = shelf_label
                incorrect['book'] = text
                # print("SC=1 incorrect",text)
                incorrect_books += 1
                misplaced.append((shelf_label, text))
                wrong = False

                df = df.append(incorrect, ignore_index=True)

        total_inc += incorrect_books
    # print(df)
    print("Correct = ", correct, "incorrect per file = ", incorrect_books)
    print("total inc = ", total_inc)

    print(misplaced, len(misplaced))
    time.sleep(1)

    document.add_paragraph('\n')
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Shelf'
    hdr_cells[1].text = 'Book Label'

    for item in misplaced:
    	print(item[0], item[1])
    	s = item[0]
    	l = item[1]
    	row_cells = table.add_row().cells
    	row_cells[0].text = s
    	row_cells[1].text = l

    p = document.add_paragraph('\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_name, width = Inches(6))

    

    # writer = pd.ExcelWriter('mismatch.xlsx', engine='openpyxl')

    # df.to_excel(writer, sheet_name = 'Report')
    # workbook = writer.book
    # worksheet = writer.sheets['Report']
    # img = oim(img_name)
    # img.height = 224
    # img.width = 224
    # worksheet.add_image(img,'E2')
    # writer.save()

    df.to_csv("mismatch.csv", mode='a', header=False)
